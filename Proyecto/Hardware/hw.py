# -*- coding: utf-8 -*-
"""
Created on Fri May 24 06:48:44 2024

@author: Maxwell
"""

from docx import Document

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading('Comparativa de Sensores y Chipsets de Alimentación', level=1)

# Agregar la tabla de comparación de sensores
doc.add_heading('Comparativa de Sensores', level=2)

# Agregar una tabla al documento para comparar sensores
sensor_table = doc.add_table(rows=1, cols=7)
sensor_table.style = 'Table Grid'

# Definir los encabezados de la tabla de sensores
sensor_hdr_cells = sensor_table.rows[0].cells
sensor_hdr_cells[0].text = 'Sensor'
sensor_hdr_cells[1].text = 'Tipo'
sensor_hdr_cells[2].text = 'Interfaz'
sensor_hdr_cells[3].text = 'Rango'
sensor_hdr_cells[4].text = 'Precisión'
sensor_hdr_cells[5].text = 'Consumo'
sensor_hdr_cells[6].text = 'Características Adicionales'

# Datos de los sensores
sensor_data = [
    ['LSM9DS1', 'IMU', 'I2C, SPI', 'Aceleración: ±2/±4/±8/±16 g, Giroscopio: ±245/±500/±2000 dps, Magnetómetro: ±4/±8/±12/±16 gauss', 'Acelerómetro: 0.00061 g/LSB, Giroscopio: 0.00875 dps/LSB, Magnetómetro: 0.00014 gauss/LSB', '2 mA', 'Acelerómetro, Giroscopio, Magnetómetro'],
    ['BNO055', 'IMU', 'I2C, UART', 'Aceleración: ±2/±4/±8/±16 g, Giroscopio: ±125/±250/±500/±1000/±2000 dps, Magnetómetro: ±1300 µT', 'Acelerómetro: 0.01 g/LSB, Giroscopio: 3.8 mDPS/LSB, Magnetómetro: 0.3 µT/LSB', '12 mA', 'Acelerómetro, Giroscopio, Magnetómetro, Fusión de sensores'],
    ['MPU6050', 'IMU', 'I2C', 'Aceleración: ±2/±4/±8/±16 g, Giroscopio: ±250/±500/±1000/±2000 dps', 'Acelerómetro: 0.000598 g/LSB, Giroscopio: 0.00763 dps/LSB', '3.9 mA', 'Acelerómetro, Giroscopio'],
    ['HC-SR04', 'Ultrasonido', 'GPIO', '2 cm - 400 cm', '±3 mm', '15 mA', 'Sensor de distancia ultrasonido']
]

# Agregar los datos de los sensores a la tabla
for sensor_row_data in sensor_data:
    sensor_row_cells = sensor_table.add_row().cells
    for i, sensor_cell_data in enumerate(sensor_row_data):
        sensor_row_cells[i].text = sensor_cell_data

# Agregar la tabla de comparación de chipsets de alimentación
doc.add_heading('Comparativa de Chipsets de Alimentación', level=2)

# Agregar una tabla al documento para comparar chipsets de alimentación
power_chipset_table = doc.add_table(rows=1, cols=7)
power_chipset_table.style = 'Table Grid'

# Definir los encabezados de la tabla de chipsets de alimentación
power_chipset_hdr_cells = power_chipset_table.rows[0].cells
power_chipset_hdr_cells[0].text = 'Chipset'
power_chipset_hdr_cells[1].text = 'Voltaje de Entrada'
power_chipset_hdr_cells[2].text = 'Voltaje de Salida'
power_chipset_hdr_cells[3].text = 'Corriente de Carga Máxima'
power_chipset_hdr_cells[4].text = 'Eficiencia'
power_chipset_hdr_cells[5].text = 'Protección de Temperatura'
power_chipset_hdr_cells[6].text = 'Desconexión Automática'

# Datos de los chipsets de alimentación
power_chipset_data = [
    ['Texas Instruments BQ24075', '4.5V - 6.5V', '4.2V', '1A', '90%', 'Sí', 'Sí'],
    ['Microchip MCP73831', '3.75V - 6V', '4.2V', '500mA', '85%', 'No', 'No'],
    ['Analog Devices LTC4060', '4.5V - 10V', '4.2V', '1A', '90%', 'Sí', 'Sí'],
    ['Maxim Integrated MAX1555', '3.7V - 7V', '4.2V', '280mA', '85%', 'No', 'No'],
    ['Linear Technology LTC4054', '4.25V - 6.5V', '4.2V', '800mA', '85%', 'Sí', 'No']
]

# Agregar los datos de los chipsets de alimentación a la tabla
for power_chipset_row_data in power_chipset_data:
    power_chipset_row_cells = power_chipset_table.add_row().cells
    for i, power_chipset_cell_data in enumerate(power_chipset_row_data):
        power_chipset_row_cells[i].text = power_chipset_cell_data

# Guardar el documento
doc.save('Comparativas.docx')



